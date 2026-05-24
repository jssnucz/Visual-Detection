import gradio as gr
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import cv2
import base64
import io
import json
import re
import requests
from pathlib import Path
import easyocr
from ultralytics import YOLO
import torch
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
from datetime import datetime
import tempfile
import os

os.environ.pop('HTTP_PROXY', None)
os.environ.pop('HTTPS_PROXY', None)
os.environ.pop('http_proxy', None)
os.environ.pop('https_proxy', None)

PDF_FONT = 'Helvetica'
CHINESE_FONT_PATHS = [
    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
    '/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc',
    '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
    '/usr/share/fonts/truetype/arphic/uming.ttc',
    '/usr/share/fonts/truetype/arphic/ukai.ttc',
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    '/System/Library/Fonts/PingFang.ttc',
    '/System/Library/Fonts/STHeiti Light.ttc',
    'C:/Windows/Fonts/simhei.ttf',
    'C:/Windows/Fonts/msyh.ttc',
]

for font_path in CHINESE_FONT_PATHS:
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
            PDF_FONT = 'ChineseFont'
            break
        except:
            continue

if PDF_FONT == 'Helvetica':
    try:
        import subprocess
        font_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')
        os.makedirs(font_dir, exist_ok=True)
        font_file = os.path.join(font_dir, 'NotoSansSC-Regular.ttf')
        if not os.path.exists(font_file):
            subprocess.run([
                'wget', '-q', '-O', font_file,
                'https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf'
            ], timeout=30)
        if os.path.exists(font_file):
            pdfmetrics.registerFont(TTFont('ChineseFont', font_file))
            PDF_FONT = 'ChineseFont'
    except:
        pass


class VisionDetectionSystem:
    def __init__(self):
        self.yolo_model = None
        self.ocr_reader = None
        self.api_key = ""
        self.api_base = "https://oa.api2d.net"
        self.model_name = "gpt-4o"
        self.last_detection_image = None
        self.last_detection_result = None
        self.last_detection_request = None
        self.load_models()
    
    def load_models(self):
        try:
            self.yolo_model = YOLO('yolov8n.pt')
        except:
            pass
        try:
            self.ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=torch.cuda.is_available())
        except:
            pass
    
    def set_api_config(self, api_key, api_base, model_name):
        self.api_key = api_key
        self.api_base = api_base if api_base else "https://api.openai.com/v1"
        self.model_name = model_name if model_name else "gpt-4o"
        return "API配置已更新"
    
    def encode_image(self, image):
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    def call_llm(self, prompt, image=None):
        if not self.api_key:
            return {"error": "请先配置API密钥"}
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        messages = []
        if image is not None:
            base64_image = self.encode_image(image)
            messages.append({
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}}
                ]
            })
        else:
            messages.append({"role": "user", "content": prompt})
        
        payload = {
            "model": self.model_name,
            "messages": messages,
            "max_tokens": 2000
        }
        
        try:
            response = requests.post(
                f"{self.api_base}/chat/completions",
                headers=headers,
                json=payload,
                timeout=60,
                proxies={"http": None, "https": None}
            )
            result = response.json()
            if "choices" in result:
                return {"content": result["choices"][0]["message"]["content"]}
            return {"error": str(result)}
        except Exception as e:
            return {"error": str(e)}
    
    def parse_detection_request(self, user_input):
        prompt = f"""用户想要检测图像中的目标。请分析用户的需求并返回JSON格式的响应。

用户输入: {user_input}

请返回以下JSON格式（不要包含其他文字）:
{{
    "targets": ["目标1", "目标2"],
    "description": "检测任务描述",
    "use_yolo": true/false,
    "yolo_classes": ["person", "car", "dog", "cat", "bird", "bicycle", "motorcycle", "bus", "truck", "traffic light", "fire hydrant", "stop sign", "bench", "elephant", "bear", "zebra", "giraffe", "backpack", "umbrella", "handbag", "tie", "suitcase", "frisbee", "skis", "snowboard", "sports ball", "kite", "baseball bat", "baseball glove", "skateboard", "surfboard", "tennis racket", "bottle", "wine glass", "cup", "fork", "knife", "spoon", "bowl", "banana", "apple", "sandwich", "orange", "broccoli", "carrot", "hot dog", "pizza", "donut", "cake", "chair", "couch", "potted plant", "bed", "dining table", "toilet", "tv", "laptop", "mouse", "remote", "keyboard", "cell phone", "microwave", "oven", "toaster", "sink", "refrigerator", "book", "clock", "vase", "scissors", "teddy bear", "hair drier", "toothbrush"],
    "custom_prompt": "用于视觉大模型的检测提示"
}}

如果用户要检测的目标在YOLO可检测类别中，设置use_yolo为true并指定相应类别。
否则设置use_yolo为false，并生成custom_prompt用于视觉大模型检测。"""

        result = self.call_llm(prompt)
        if "error" in result:
            return None, result["error"]
        
        try:
            content = result["content"]
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                return json.loads(json_match.group()), None
        except:
            pass
        return None, "解析失败"
    
    def detect_with_yolo(self, image, classes=None):
        if self.yolo_model is None:
            return image, "YOLO模型未加载"
        
        if isinstance(image, Image.Image):
            image = np.array(image)
        
        results = self.yolo_model(image)
        annotated = image.copy()
        
        detections = []
        for r in results:
            boxes = r.boxes
            for box in boxes:
                cls = int(box.cls[0])
                conf = float(box.conf[0])
                name = self.yolo_model.names[cls]
                
                if classes and name.lower() not in [c.lower() for c in classes]:
                    continue
                
                x1, y1, x2, y2 = map(int, box.xyxy[0])
                detections.append({
                    "class": name,
                    "confidence": conf,
                    "bbox": [x1, y1, x2, y2]
                })
                
                cv2.rectangle(annotated, (x1, y1), (x2, y2), (0, 255, 0), 3)
                label = f"{name}: {conf:.2f}"
                cv2.putText(annotated, label, (x1, y1-10), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (0, 255, 0), 2)
        
        return Image.fromarray(annotated), detections
    
    def detect_with_vision_model(self, image, custom_prompt):
        prompt = f"""{custom_prompt}

请分析图像并返回检测到的目标位置。使用以下JSON格式返回（不要包含其他文字）:
{{
    "detections": [
        {{
            "object": "目标名称",
            "location": "位置描述（如：左上角、中间、右下等）",
            "bbox_percent": [x1%, y1%, x2%, y2%]
        }}
    ],
    "description": "整体描述"
}}

bbox_percent使用百分比表示边界框位置，范围0-100。"""

        result = self.call_llm(prompt, image)
        if "error" in result:
            return image, result["error"]
        
        try:
            content = result["content"]
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                data = json.loads(json_match.group())
                return self.draw_detections(image, data)
        except Exception as e:
            return image, f"解析错误: {str(e)}"
        
        return image, content
    
    def draw_detections(self, image, data):
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        
        draw = ImageDraw.Draw(image)
        width, height = image.size
        
        try:
            font = ImageFont.truetype("arial.ttf", 20)
        except:
            font = ImageFont.load_default()
        
        detections = data.get("detections", [])
        colors = [(255, 0, 0), (0, 255, 0), (0, 0, 255), (255, 255, 0), (255, 0, 255), (0, 255, 255)]
        
        for i, det in enumerate(detections):
            color = colors[i % len(colors)]
            bbox = det.get("bbox_percent", [])
            if len(bbox) == 4:
                x1 = int(bbox[0] * width / 100)
                y1 = int(bbox[1] * height / 100)
                x2 = int(bbox[2] * width / 100)
                y2 = int(bbox[3] * height / 100)
                
                draw.rectangle([x1, y1, x2, y2], outline=color, width=3)
                label = det.get("object", "Unknown")
                draw.text((x1, y1-25), label, fill=color, font=font)
        
        return image, data.get("description", "检测完成")
    
    def process_detection(self, image, user_input):
        if image is None:
            return None, "请上传图片"
        
        if not user_input:
            return image, "请输入检测需求"
        
        self.last_detection_request = user_input
        
        parsed, error = self.parse_detection_request(user_input)
        if error:
            return image, f"解析错误: {error}"
        
        if parsed.get("use_yolo", False) and self.yolo_model:
            classes = parsed.get("yolo_classes", [])
            targets = parsed.get("targets", [])
            filter_classes = [c for c in classes if any(t.lower() in c.lower() or c.lower() in t.lower() for t in targets)]
            result_image, detections = self.detect_with_yolo(image, filter_classes if filter_classes else None)
            result_text = f"YOLO检测结果:\n{json.dumps(detections, ensure_ascii=False, indent=2)}"
            self.last_detection_image = result_image
            self.last_detection_result = result_text
            return result_image, result_text
        else:
            custom_prompt = parsed.get("custom_prompt", f"请检测图像中的{', '.join(parsed.get('targets', ['目标']))}")
            result_image, result_text = self.detect_with_vision_model(image, custom_prompt)
            self.last_detection_image = result_image
            self.last_detection_result = result_text
            return result_image, result_text
    
    def export_to_docx(self, image, result_text, request_text):
        if image is None or result_text is None:
            return None
        
        doc = DocxDocument()
        
        title = doc.add_heading('智能目标检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        doc.add_heading('检测需求', level=1)
        doc.add_paragraph(request_text if request_text else "未指定")
        
        doc.add_heading('检测结果图像', level=1)
        
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        doc.add_picture(img_buffer, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('检测详情', level=1)
        
        try:
            if "YOLO检测结果" in result_text:
                json_str = result_text.replace("YOLO检测结果:\n", "")
                detections = json.loads(json_str)
                
                if detections:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '序号'
                    hdr_cells[1].text = '目标类别'
                    hdr_cells[2].text = '置信度'
                    hdr_cells[3].text = '边界框坐标'
                    
                    for cell in hdr_cells:
                        cell.paragraphs[0].runs[0].bold = True
                    
                    for idx, det in enumerate(detections, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(idx)
                        row_cells[1].text = det.get('class', 'N/A')
                        row_cells[2].text = f"{det.get('confidence', 0):.2%}"
                        bbox = det.get('bbox', [])
                        row_cells[3].text = f"[{bbox[0]}, {bbox[1]}, {bbox[2]}, {bbox[3]}]" if len(bbox) == 4 else "N/A"
                else:
                    doc.add_paragraph("未检测到目标")
            else:
                doc.add_paragraph(result_text)
        except:
            doc.add_paragraph(result_text)
        
        doc.add_paragraph('')
        # doc.add_paragraph('--- 报告结束 ---').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    def export_to_pdf(self, image, result_text, request_text):
        if image is None or result_text is None:
            return None
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            pdf_path = tmp.name
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=50, bottomMargin=50)
        styles = getSampleStyleSheet()
        
        try:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontName=PDF_FONT,
                fontSize=24,
                alignment=1,
                spaceAfter=20
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontName=PDF_FONT,
                fontSize=16,
                spaceAfter=12
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=PDF_FONT,
                fontSize=12,
                spaceAfter=8
            )
        except:
            title_style = styles['Title']
            heading_style = styles['Heading1']
            normal_style = styles['Normal']
        
        story = []
        
        story.append(Paragraph('智能目标检测报告', title_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', normal_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph('检测需求', heading_style))
        story.append(Paragraph(request_text if request_text else "未指定", normal_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph('检测结果图像', heading_style))
        
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        img_width = 5 * inch
        img_height = img_width * image.height / image.width
        if img_height > 6 * inch:
            img_height = 6 * inch
            img_width = img_height * image.width / image.height
        
        story.append(RLImage(img_buffer, width=img_width, height=img_height))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph('检测详情', heading_style))
        
        try:
            if "YOLO检测结果" in result_text:
                json_str = result_text.replace("YOLO检测结果:\n", "")
                detections = json.loads(json_str)
                
                if detections:
                    table_data = [['序号', '目标类别', '置信度', '边界框坐标']]
                    for idx, det in enumerate(detections, 1):
                        bbox = det.get('bbox', [])
                        bbox_str = f"[{bbox[0]}, {bbox[1]}, {bbox[2]}, {bbox[3]}]" if len(bbox) == 4 else "N/A"
                        table_data.append([
                            str(idx),
                            det.get('class', 'N/A'),
                            f"{det.get('confidence', 0):.2%}",
                            bbox_str
                        ])
                    
                    table = Table(table_data, colWidths=[0.6*inch, 1.5*inch, 1*inch, 2.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), PDF_FONT),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('FONTNAME', (0, 1), (-1, -1), PDF_FONT),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                else:
                    story.append(Paragraph("未检测到目标", normal_style))
            else:
                for line in result_text.split('\n'):
                    if line.strip():
                        story.append(Paragraph(line, normal_style))
        except:
            for line in result_text.split('\n'):
                if line.strip():
                    try:
                        story.append(Paragraph(line, normal_style))
                    except:
                        story.append(Paragraph(line.encode('utf-8', errors='ignore').decode('utf-8'), normal_style))
        
        story.append(Spacer(1, 30))
        # story.append(Paragraph('--- 报告结束 ---', ParagraphStyle('Center', parent=normal_style, alignment=1)))
        
        doc.build(story)
        return pdf_path
    
    def do_export(self, export_format, detect_image, detect_result, detect_request):
        if detect_image is None:
            return None
        
        if export_format == "DOCX":
            return self.export_to_docx(detect_image, detect_result, detect_request)
        else:
            return self.export_to_pdf(detect_image, detect_result, detect_request)
    
    def extract_text(self, image):
        if image is None:
            return None, "请上传图片"
        
        if self.ocr_reader is None:
            if self.api_key:
                prompt = "请提取图像中的所有文字内容，保持原有的排版格式。"
                result = self.call_llm(prompt, image)
                if "error" in result:
                    return image, result["error"]
                return image, result["content"]
            return image, "OCR模块未加载且API未配置"
        
        if isinstance(image, Image.Image):
            image_np = np.array(image)
        else:
            image_np = image
        
        results = self.ocr_reader.readtext(image_np)
        
        annotated = image_np.copy()
        extracted_text = []
        
        for (bbox, text, conf) in results:
            pts = np.array(bbox, dtype=np.int32)
            cv2.polylines(annotated, [pts], True, (0, 255, 0), 2)
            extracted_text.append(f"{text} ")
            
            x, y = int(bbox[0][0]), int(bbox[0][1])
            cv2.putText(annotated, text, (x, y-5), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 0, 0), 2)
        
        return Image.fromarray(annotated), "\n".join(extracted_text)


system = VisionDetectionSystem()

def create_interface():
    with gr.Blocks(
        title="智能视觉检测系统",
        theme=gr.themes.Soft(
            primary_hue="blue",
            secondary_hue="cyan",
            neutral_hue="slate"
        ),
        css="""
        .gradio-container {
            max-width: 1400px !important;
            margin: auto !important;
        }
        .main-title {
            text-align: center;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-size: 2.5em;
            font-weight: bold;
            padding: 20px;
        }
        .subtitle {
            text-align: center;
            color: #666;
            margin-bottom: 30px;
        }
        .tab-nav button {
            font-size: 1.1em !important;
            padding: 15px 30px !important;
        }
        .output-box {
            min-height: 200px;
            border-radius: 10px;
        }
        """
    ) as demo:
        gr.HTML("""
            <div class="main-title">🔍 智能视觉检测系统</div>
            <div class="subtitle">集成自然语言理解 + 图像目标检测 + OCR文字识别</div>
        """)
        
        with gr.Tabs() as tabs:
            with gr.TabItem("🎯 智能目标检测", id=0):
                gr.Markdown("""
                ### 使用说明
                1. 上传需要检测的图片
                2. 用自然语言描述您想检测的目标（如："帮我找出图片中的所有人"、"检测图中的汽车"）
                3. 系统会自动选择最合适的检测方式并标注结果
                4. 检测完成后可导出结果为DOCX或PDF文件
                """)
                
                with gr.Row():
                    with gr.Column(scale=1):
                        detect_image_input = gr.Image(
                            label="📷 上传图片",
                            type="pil",
                            height=400
                        )
                        detect_text_input = gr.Textbox(
                            label="💬 输入检测需求",
                            placeholder="例如：帮我找出图片中的所有人、检测图中的猫、找出所有红色的物体...",
                            lines=2
                        )
                        detect_btn = gr.Button("🚀 开始检测", variant="primary", size="lg")
                    
                    with gr.Column(scale=1):
                        detect_image_output = gr.Image(
                            label="🖼️ 检测结果",
                            type="pil",
                            height=400
                        )
                        detect_text_output = gr.Textbox(
                            label="📋 检测详情",
                            lines=8,
                            elem_classes=["output-box"]
                        )
                        
                        with gr.Row():
                            export_format = gr.Radio(
                                choices=["DOCX", "PDF"],
                                value="DOCX",
                                label="📁 导出格式"
                            )
                            export_btn = gr.Button("📥 导出报告", variant="secondary", size="lg")
                        
                        export_file = gr.File(label="📄 下载文件")
                
                detect_btn.click(
                    fn=system.process_detection,
                    inputs=[detect_image_input, detect_text_input],
                    outputs=[detect_image_output, detect_text_output]
                )
                
                export_btn.click(
                    fn=system.do_export,
                    inputs=[export_format, detect_image_output, detect_text_output, detect_text_input],
                    outputs=[export_file]
                )
            
            with gr.TabItem("📝 OCR文字识别", id=1):
                gr.Markdown("""
                ### 使用说明
                1. 上传包含文字的图片
                2. 点击识别按钮，系统将自动提取图片中的所有文字
                3. 支持中英文混合识别
                """)
                
                with gr.Row():
                    with gr.Column(scale=1):
                        ocr_image_input = gr.Image(
                            label="📷 上传图片",
                            type="pil",
                            height=400
                        )
                        ocr_btn = gr.Button("🔤 提取文字", variant="primary", size="lg")
                    
                    with gr.Column(scale=1):
                        ocr_image_output = gr.Image(
                            label="🖼️ 标注结果",
                            type="pil",
                            height=400
                        )
                        ocr_text_output = gr.Textbox(
                            label="📄 识别文字",
                            lines=10,
                            elem_classes=["output-box"]
                        )
                
                ocr_btn.click(
                    fn=system.extract_text,
                    inputs=[ocr_image_input],
                    outputs=[ocr_image_output, ocr_text_output]
                )
            
            with gr.TabItem("⚙️ API设置", id=2):
                gr.Markdown("""
                ### API配置
                配置大语言模型API以启用高级检测功能。支持OpenAI API及兼容接口。
                """)
                
                with gr.Row():
                    with gr.Column():
                        api_key_input = gr.Textbox(
                            label="🔑 API密钥",
                            placeholder="输入您的API密钥",
                            type="password"
                        )
                        api_base_input = gr.Textbox(
                            label="🌐 API地址",
                            placeholder="https://oa.api2d.net",
                            value="https://oa.api2d.net"
                        )
                        model_input = gr.Textbox(
                            label="🤖 模型名称",
                            placeholder="gpt-4o",
                            value="gpt-4o"
                        )
                        config_btn = gr.Button("💾 保存配置", variant="primary")
                        config_status = gr.Textbox(label="状态", interactive=False)
                
                config_btn.click(
                    fn=system.set_api_config,
                    inputs=[api_key_input, api_base_input, model_input],
                    outputs=[config_status]
                )
            
            with gr.TabItem("📖 使用指南", id=3):
                gr.Markdown("""
                ## 🎯 功能介绍
                
                ### 1. 智能目标检测
                - **YOLO快速检测**：对于常见物体（人、车、动物等80类），使用YOLOv8进行快速精确检测
                - **视觉大模型检测**：对于特殊目标，调用GPT-4V等视觉大模型进行智能识别
                - **自然语言交互**：用日常语言描述您的检测需求，系统自动理解并执行
                - **结果导出**：支持将检测结果导出为DOCX或PDF格式的报告
                
                ### 2. OCR文字识别
                - 支持中英文混合识别
                - 自动标注文字位置
                - 显示识别置信度
                
                ### 3. 支持的检测目标（YOLO）
                人、自行车、汽车、摩托车、飞机、公交车、火车、卡车、船、
                交通灯、消防栓、停车标志、长凳、鸟、猫、狗、马、羊、牛、
                大象、熊、斑马、长颈鹿、背包、雨伞、手提包、领带、行李箱等80类
                
                
                ### 4. 导出功能说明
                - **DOCX格式**：生成Word文档，包含检测图像和详细结果表格
                - **PDF格式**：生成PDF文档，适合打印和分享
                
                ## 
                """)
        
        gr.HTML("""
            <div style="text-align: center; margin-top: 30px; padding: 20px; color: #888;">
                <p>智能视觉检测系统 v1.0 | Powered by YOLOv8 & GPT-4V & EasyOCR</p>
            </div>
        """)
    
    return demo


if __name__ == "__main__":
    demo = create_interface()
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=False,
        inbrowser=True
    )